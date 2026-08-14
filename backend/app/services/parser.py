# import io
# from docx import Document
# import pdfplumber


# # def parse_docx(file_bytes: bytes) -> list[str]:
# #     doc = Document(io.BytesIO(file_bytes))
# #     segments = []

# #     # walk the document body in order so tables stay near their
# #     # surrounding paragraphs instead of being appended at the end
# #     for element in doc.element.body:
# #         if element.tag.endswith('}p'):  # paragraph
# #             para = next((p for p in doc.paragraphs if p._element is element), None)
# #             if para and para.text.strip():
# #                 segments.append(para.text.strip())
# #         elif element.tag.endswith('}tbl'):  # table
# #             table = next((t for t in doc.tables if t._element is element), None)
# #             if table:
# #                 segments.append(_table_to_text(table.rows))

# #     return segments


# def parse_docx(file_bytes: bytes) -> list[str]:
#     doc = Document(io.BytesIO(file_bytes))
#     segments = []

#     # build O(1) lookup maps once, instead of scanning doc.paragraphs / doc.tables
#     # for every single element (which was O(n) per element = O(n²) overall)
#     para_by_element = {p._element: p for p in doc.paragraphs}
#     table_by_element = {t._element: t for t in doc.tables}

#     for element in doc.element.body:
#         if element.tag.endswith('}p'):
#             para = para_by_element.get(element)
#             if para and para.text.strip():
#                 segments.append(para.text.strip())
#         elif element.tag.endswith('}tbl'):
#             table = table_by_element.get(element)
#             if table:
#                 segments.append(_table_to_text(table.rows))

#     return segments


# def _table_to_text(rows) -> str:
#     """Flatten a table's rows into a readable text block."""
#     rows_text = []
#     for row in rows:
#         # docx rows have .cells; pdfplumber rows are already plain lists of strings
#         cells = [c.text.strip() for c in row.cells] if hasattr(row, "cells") else [c or "" for c in row]
#         rows_text.append(" | ".join(cells))
#     return "[TABLE]\n" + "\n".join(rows_text)


# def _in_any_bbox(x0: float, top: float, bboxes: list) -> bool:
#     for bbox in bboxes:
#         if bbox[0] - 1 <= x0 <= bbox[2] + 1 and bbox[1] - 1 <= top <= bbox[3] + 1:
#             return True
#     return False


# def parse_pdf(file_bytes: bytes) -> list[str]:
#     """
#     Extract paragraph-level text and tables from a PDF.

#     PDFs have no real paragraph markers, so breaks are inferred from the
#     vertical gap between lines, relative to each line's own text height
#     (so it scales with font size instead of using a fixed point value,
#     which doesn't generalize across PDFs from different generators).
#     """
#     GAP_RATIO = 0.5  # gap counts as a paragraph break if > 50% of the line's height

#     segments = []

#     with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
#         for page in pdf.pages:
#             table_bboxes = []
#             for table in page.find_tables():
#                 extracted = table.extract()
#                 if extracted:
#                     segments.append(_table_to_text(extracted))
#                     table_bboxes.append(table.bbox)

#             lines = page.extract_text_lines()
#             current_para = None
#             prev_bottom = None

#             for line in lines:
#                 if _in_any_bbox(line['x0'], line['top'], table_bboxes):
#                     continue

#                 text = line['text'].strip()
#                 if not text:
#                     continue

#                 line_height = line['bottom'] - line['top']
#                 gap = (line['top'] - prev_bottom) if prev_bottom is not None else None
#                 is_break = gap is None or gap > (line_height * GAP_RATIO)

#                 if current_para is not None and not is_break:
#                     current_para += " " + text
#                 else:
#                     if current_para is not None:
#                         segments.append(current_para)
#                     current_para = text

#                 prev_bottom = line['bottom']

#             if current_para is not None:
#                 segments.append(current_para)

#     return segments


# def parse_document(file_bytes: bytes, filename: str) -> list[str]:
#     """
#     Dispatches to the correct parser based on file extension.
#     Raises ValueError for unsupported types.
#     """
#     ext = filename.rsplit(".", 1)[-1].lower()

#     if ext == "docx":
#         return parse_docx(file_bytes)
#     elif ext == "pdf":
#         return parse_pdf(file_bytes)
#     else:
#         raise ValueError(f"Unsupported file type: {ext}")













import io
import zipfile
from docx import Document
import pdfplumber
from lxml import etree
from app.services.numbering import load_numbering_defs, NumberingTracker
from app.services.comments import load_comments, get_comment_ids_for_paragraph

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def _qn(tag):
    return f"{{{W_NS}}}{tag}"

# def _extract_paragraph(p_element, tracker: NumberingTracker | None) -> dict:
#     text_parts = []
#     authors = set()

#     for el in p_element.iter():
#         tag = etree.QName(el).localname
#         if tag == "t":
#             parent = el.getparent()
#             inside_del = False
#             while parent is not None and parent is not p_element:
#                 if etree.QName(parent).localname == "del":
#                     inside_del = True
#                     break
#                 parent = parent.getparent()
#             if not inside_del:
#                 text_parts.append(el.text or "")
#         elif tag in ("ins", "del"):
#             author = el.get(_qn("author"))
#             if author:
#                 authors.add(author)

#     text = "".join(text_parts).strip()

#     para_mark_deleted = False
#     pPr = p_element.find("w:pPr", NS)
#     if pPr is not None:
#         rPr = pPr.find("w:rPr", NS)
#         if rPr is not None and rPr.find("w:del", NS) is not None:
#             para_mark_deleted = True

#     rendered_number = None
#     if tracker is not None and not para_mark_deleted:
#         numPr = p_element.find(".//w:numPr", NS)
#         if numPr is not None:
#             numId_el = numPr.find("w:numId", NS)
#             ilvl_el = numPr.find("w:ilvl", NS)
#             numId = numId_el.get(_qn("val")) if numId_el is not None else None
#             ilvl = ilvl_el.get(_qn("val")) if ilvl_el is not None else "0"
#             if numId is not None:
#                 rendered_number = tracker.render(numId, ilvl)

#     if rendered_number and text:
#         text = f"{rendered_number} {text}"

#     return {"text": text, "authors": sorted(authors)}

def _extract_paragraph(p_element, tracker: NumberingTracker | None, mode: str = "accepted") -> dict:
    text_parts = []
    authors = set()

    for el in p_element.iter():
        tag = etree.QName(el).localname

        if tag == "t":
            parent = el.getparent()
            inside_del = False
            inside_ins = False
            check = parent
            while check is not None and check is not p_element:
                lname = etree.QName(check).localname
                if lname == "del":
                    inside_del = True
                if lname == "ins":
                    inside_ins = True
                check = check.getparent()

            if mode == "accepted":
                if not inside_del:
                    text_parts.append(el.text or "")
            else:
                if not inside_ins:
                    text_parts.append(el.text or "")

        elif tag == "delText":
            if mode == "original":
                text_parts.append(el.text or "")

        elif tag in ("ins", "del"):
            author = el.get(_qn("author"))
            if author:
                authors.add(author)

    text = "".join(text_parts).strip()

    # <<< NEW BLOCK GOES HERE >>>
    # If the paragraph's own MARK (not just its wording) was deleted,
    # Word does not treat it as visible content even in "Original"/
    # reject-changes view — confirmed against real Word and an
    # independent reconstruction. Suppress its text here too, not just
    # its numbering slot, so we don't surface content that never
    # actually appeared to a human reader.
    mark_deleted = False
    pPr = p_element.find("w:pPr", NS)
    if pPr is not None:
        rPr = pPr.find("w:rPr", NS)
        if rPr is not None and rPr.find("w:del", NS) is not None:
            mark_deleted = True

    if mode == "original" and mark_deleted:
        text = ""
    # <<< END NEW BLOCK >>>

    rendered_number = None
    if tracker is not None and text:
        numPr = p_element.find(".//w:numPr", NS)
        if numPr is not None:
            numId_el = numPr.find("w:numId", NS)
            ilvl_el = numPr.find("w:ilvl", NS)
            numId = numId_el.get(_qn("val")) if numId_el is not None else None
            ilvl = ilvl_el.get(_qn("val")) if ilvl_el is not None else "0"
            if numId is not None:
                rendered_number = tracker.render(numId, ilvl)

    if rendered_number and text:
        text = f"{rendered_number} {text}"

    return {"text": text, "authors": sorted(authors)}


CELL_LINE_BREAK = "\u2028"  # distinct from the \n used to separate rows

def _extract_table(tbl_element, tracker: NumberingTracker | None = None, mode: str = "accepted") -> dict:
    rows_out = []
    authors = set()

    for tr in tbl_element.findall(".//w:tr", NS):
        cells_out = []
        for tc in tr.findall("w:tc", NS):
            paragraph_texts = []
            for p in tc.findall("w:p", NS):
                data = _extract_paragraph(p, tracker, mode)
                if data["text"]:
                    paragraph_texts.append(data["text"])
                authors.update(data["authors"])
            cells_out.append(CELL_LINE_BREAK.join(paragraph_texts))
        rows_out.append(cells_out)

    rows_text = [" | ".join(cells) for cells in rows_out]
    return {"text": "[TABLE]\n" + "\n".join(rows_text), "authors": sorted(authors)}


def parse_docx(file_bytes: bytes, mode: str = "accepted") -> list[dict]:
    with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
        xml_bytes = z.read("word/document.xml")
    root = etree.fromstring(xml_bytes)
    body = root.find(".//w:body", NS)

    num_to_abstract, abstract_levels = load_numbering_defs(file_bytes)
    tracker = NumberingTracker(num_to_abstract, abstract_levels) if num_to_abstract else None
    comments_lookup = load_comments(file_bytes)

    segments = []
    for element in body:
        tag = etree.QName(element).localname
        if tag == "p":
            data = _extract_paragraph(element, tracker, mode)
            if data["text"]:
                comment_ids = get_comment_ids_for_paragraph(element)
                data["comments"] = [comments_lookup[cid] for cid in comment_ids if cid in comments_lookup]
                segments.append(data)
        elif tag == "tbl":
            data = _extract_table(element, tracker, mode)
            if data["text"].strip() != "[TABLE]":
                data["comments"] = []
                segments.append(data)

    return segments


def parse_document(file_bytes: bytes, filename: str, mode: str = "accepted") -> list[dict]:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "docx":
        return parse_docx(file_bytes, mode)
    elif ext == "pdf":
        return parse_pdf(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _table_to_text(rows) -> str:
    rows_text = []
    for row in rows:
        cells = [c.text.strip() for c in row.cells] if hasattr(row, "cells") else [c or "" for c in row]
        rows_text.append(" | ".join(cells))
    return "[TABLE]\n" + "\n".join(rows_text)


def _in_any_bbox(x0: float, top: float, bboxes: list) -> bool:
    for bbox in bboxes:
        if bbox[0] - 1 <= x0 <= bbox[2] + 1 and bbox[1] - 1 <= top <= bbox[3] + 1:
            return True
    return False


def parse_pdf(file_bytes: bytes) -> list[dict]:
    # unchanged from your current version — PDFs can't carry numbering
    # definitions the same way, out of scope here
    GAP_RATIO = 0.5
    segments = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            table_bboxes = []
            for table in page.find_tables():
                extracted = table.extract()
                if extracted:
                    segments.append({"text": _table_to_text(extracted), "authors": []})
                    table_bboxes.append(table.bbox)

            lines = page.extract_text_lines()
            current_para = None
            prev_bottom = None

            for line in lines:
                if _in_any_bbox(line['x0'], line['top'], table_bboxes):
                    continue
                text = line['text'].strip()
                if not text:
                    continue
                line_height = line['bottom'] - line['top']
                gap = (line['top'] - prev_bottom) if prev_bottom is not None else None
                is_break = gap is None or gap > (line_height * GAP_RATIO)

                if current_para is not None and not is_break:
                    current_para += " " + text
                else:
                    if current_para is not None:
                        segments.append({"text": current_para, "authors": []})
                    current_para = text
                prev_bottom = line['bottom']

            if current_para is not None:
                segments.append({"text": current_para, "authors": []})

    return segments


# def parse_document(file_bytes: bytes, filename: str) -> list[dict]:
#     ext = filename.rsplit(".", 1)[-1].lower()
#     if ext == "docx":
#         return parse_docx(file_bytes)
#     elif ext == "pdf":
#         return parse_pdf(file_bytes)
#     else:
#         raise ValueError(f"Unsupported file type: {ext}")