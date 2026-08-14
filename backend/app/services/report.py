# import io
# from docx import Document
# from docx.shared import RGBColor

# RISK_COLORS = {
#     "high": RGBColor(0xC1, 0x27, 0x2D),
#     "medium": RGBColor(0xC9, 0x9A, 0x2E),
#     "low": RGBColor(0x2E, 0x7D, 0xC9),
#     "cosmetic": RGBColor(0x6B, 0x72, 0x80),
# }


# def generate_report(comparison_id: str, original_filename: str, revised_filename: str, clauses: list[dict]) -> bytes:
#     doc = Document()

#     doc.add_heading("Contract Comparison Summary", level=0)

#     meta = doc.add_paragraph()
#     meta.add_run(f"Original: {original_filename}\n").bold = True
#     meta.add_run(f"Revised: {revised_filename}\n").bold = True

#     changed = [c for c in clauses if c["status"] != "unchanged"]
#     summary = doc.add_paragraph()
#     summary.add_run(f"{len(changed)} change(s) found out of {len(clauses)} sections reviewed.").italic = True

#     doc.add_paragraph()

#     for c in changed:
#         heading = doc.add_paragraph()
#         status_run = heading.add_run(f"[{c['status'].upper()}] ")
#         status_run.bold = True

#         if c.get("risk_level"):
#             risk_run = heading.add_run(f"{c['risk_level'].upper()} RISK")
#             risk_run.bold = True
#             risk_run.font.color.rgb = RISK_COLORS.get(c["risk_level"], RGBColor(0, 0, 0))

#         if c.get("ai_summary"):
#             summary_p = doc.add_paragraph()
#             summary_p.add_run(c["ai_summary"]).italic = True

#         if c.get("authors"):
#             authors_p = doc.add_paragraph()
#             authors_run = authors_p.add_run(f"Edited by: {', '.join(c['authors'])}")
#             authors_run.italic = True
#             authors_run.font.size = authors_run.font.size or None
#             authors_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

#         if c["status"] == "modified":
#             orig_p = doc.add_paragraph()
#             orig_p.add_run("Original: ").bold = True
#             orig_run = orig_p.add_run(c["original_text"] or "")
#             orig_run.font.strike = True
#             orig_run.font.color.rgb = RGBColor(0xB2, 0x3A, 0x2E)

#             rev_p = doc.add_paragraph()
#             rev_p.add_run("Revised: ").bold = True
#             rev_run = rev_p.add_run(c["revised_text"] or "")
#             rev_run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x5C)

#         elif c["status"] == "added":
#             p = doc.add_paragraph()
#             p.add_run("Added: ").bold = True
#             p.add_run(c["revised_text"] or "")

#         elif c["status"] == "deleted":
#             p = doc.add_paragraph()
#             p.add_run("Removed: ").bold = True
#             del_run = p.add_run(c["original_text"] or "")
#             del_run.font.strike = True

#         doc.add_paragraph()  # spacing between clauses

#     buffer = io.BytesIO()
#     doc.save(buffer)
#     return buffer.getvalue()








import io
from docx import Document
from docx.shared import RGBColor

RISK_COLORS = {
    "high": RGBColor(0xC1, 0x27, 0x2D),
    "medium": RGBColor(0xC9, 0x9A, 0x2E),
    "low": RGBColor(0x2E, 0x7D, 0xC9),
    "cosmetic": RGBColor(0x6B, 0x72, 0x80),
}

CELL_LINE_BREAK = "\u2028"  # matches the marker used in parser.py's table extraction


def _add_table_or_text(doc, label: str, text: str, strike: bool = False, color: RGBColor = None):
    """
    Renders table-formatted text ("[TABLE]\\n row | row | row") as a real
    Word table with proper rows/columns, since dumping that flattened
    string into one plain paragraph made it unreadable. Cells that contain
    multiple lines (joined with CELL_LINE_BREAK, not \\n, so they don't get
    mistaken for row separators) render as " / "-joined text within the cell.
    Everything else still renders as a plain labeled paragraph, with
    optional strike/color formatting.
    """
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True

    if text.startswith("[TABLE]"):
        rows_raw = text.replace("[TABLE]\n", "").split("\n")
        rows = [row.split(" | ") for row in rows_raw if row.strip()]
        if rows:
            col_count = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=col_count)
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci in range(col_count):
                    cell_text = row[ci] if ci < len(row) else ""
                    table.rows[ri].cells[ci].text = cell_text.replace(CELL_LINE_BREAK, " / ")
    else:
        run = p.add_run(text)
        if strike:
            run.font.strike = True
        if color:
            run.font.color.rgb = color


def generate_report(comparison_id: str, original_filename: str, revised_filename: str, clauses: list[dict]) -> bytes:
    doc = Document()

    doc.add_heading("Contract Comparison Summary", level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"Original: {original_filename}\n").bold = True
    meta.add_run(f"Revised: {revised_filename}\n").bold = True

    changed = [c for c in clauses if c["status"] != "unchanged"]
    summary = doc.add_paragraph()
    summary.add_run(f"{len(changed)} change(s) found out of {len(clauses)} sections reviewed.").italic = True

    doc.add_paragraph()

    for c in changed:
        heading = doc.add_paragraph()
        status_run = heading.add_run(f"[{c['status'].upper()}] ")
        status_run.bold = True

        if c.get("risk_level"):
            risk_run = heading.add_run(f"{c['risk_level'].upper()} RISK")
            risk_run.bold = True
            risk_run.font.color.rgb = RISK_COLORS.get(c["risk_level"], RGBColor(0, 0, 0))

        if c.get("ai_summary"):
            summary_p = doc.add_paragraph()
            summary_p.add_run(c["ai_summary"]).italic = True

        if c.get("authors"):
            authors_p = doc.add_paragraph()
            authors_run = authors_p.add_run(f"Edited by: {', '.join(c['authors'])}")
            authors_run.italic = True
            authors_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        if c.get("comments"):
            for cm in c["comments"]:
                comment_p = doc.add_paragraph()
                comment_run = comment_p.add_run(f"\U0001F4AC {cm['author']}: {cm['text']}")
                comment_run.italic = True
                comment_run.font.color.rgb = RGBColor(0x8B, 0x93, 0xA6)

        if c["status"] == "modified":
            _add_table_or_text(doc, "Original", c["original_text"] or "", strike=True, color=RGBColor(0xB2, 0x3A, 0x2E))
            _add_table_or_text(doc, "Revised", c["revised_text"] or "", color=RGBColor(0x1F, 0x7A, 0x5C))

        elif c["status"] == "added":
            _add_table_or_text(doc, "Added", c["revised_text"] or "")

        elif c["status"] == "deleted":
            _add_table_or_text(doc, "Removed", c["original_text"] or "", strike=True)

        doc.add_paragraph()  # spacing between clauses

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()